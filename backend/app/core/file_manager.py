"""
File management utilities for the Academic Repository System
"""
import os
import shutil
import mimetypes
from pathlib import Path
from typing import Optional, Dict, Any, List
import magic
import hashlib
from PIL import Image, ImageOps
import fitz  # PyMuPDF
import docx
from datetime import datetime
import asyncio
import aiofiles

class FileManager:
    def __init__(self, base_upload_dir: str = "uploads"):
        self.base_dir = Path(base_upload_dir)
        self.documents_dir = self.base_dir / "documents"
        self.thumbnails_dir = self.base_dir / "thumbnails"
        self.previews_dir = self.base_dir / "previews"
        self.temp_dir = self.base_dir / "temp"
        
        # Create directories
        for directory in [self.documents_dir, self.thumbnails_dir, self.previews_dir, self.temp_dir]:
            directory.mkdir(parents=True, exist_ok=True)
    
    def get_file_info(self, file_path: Path) -> Dict[str, Any]:
        """Get comprehensive file information"""
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        stat = file_path.stat()
        mime_type, _ = mimetypes.guess_type(str(file_path))
        
        # Use python-magic for more accurate MIME type detection
        try:
            detected_mime = magic.from_file(str(file_path), mime=True)
            if detected_mime:
                mime_type = detected_mime
        except:
            pass
        
        return {
            "filename": file_path.name,
            "size": stat.st_size,
            "size_mb": round(stat.st_size / (1024 * 1024), 2),
            "mime_type": mime_type,
            "extension": file_path.suffix.lower(),
            "created": datetime.fromtimestamp(stat.st_ctime),
            "modified": datetime.fromtimestamp(stat.st_mtime),
            "md5_hash": self.calculate_file_hash(file_path)
        }
    
    def calculate_file_hash(self, file_path: Path) -> str:
        """Calculate MD5 hash of file for integrity checking"""
        hash_md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()
    
    def validate_file_type(self, file_path: Path, allowed_extensions: List[str]) -> bool:
        """Validate if file type is allowed"""
        return file_path.suffix.lower() in allowed_extensions
    
    def validate_file_size(self, file_path: Path, max_size_mb: int) -> bool:
        """Validate file size"""
        size_mb = file_path.stat().st_size / (1024 * 1024)
        return size_mb <= max_size_mb
    
    def scan_for_viruses(self, file_path: Path) -> bool:
        """Basic file scanning (placeholder for antivirus integration)"""
        # This would integrate with ClamAV or similar
        # For now, just check file size and basic patterns
        
        # Check for suspiciously large files
        if file_path.stat().st_size > 100 * 1024 * 1024:  # 100MB
            return False
        
        # Check for executable extensions in disguise
        suspicious_patterns = [b'MZ', b'PK\x03\x04', b'\x7fELF']
        try:
            with open(file_path, 'rb') as f:
                header = f.read(1024)
                for pattern in suspicious_patterns:
                    if pattern in header and file_path.suffix.lower() not in ['.zip', '.exe', '.dll']:
                        return False
        except:
            pass
        
        return True
    
    async def generate_thumbnail(self, file_path: Path, document_id: str, size: tuple = (300, 300)) -> Optional[str]:
        """Generate thumbnail for supported file types"""
        try:
            thumbnail_path = self.thumbnails_dir / f"{document_id}.jpg"
            extension = file_path.suffix.lower()
            
            if extension in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.webp']:
                # Image thumbnail
                with Image.open(file_path) as img:
                    # Convert to RGB if necessary
                    if img.mode in ('RGBA', 'LA', 'P'):
                        img = img.convert('RGB')
                    
                    # Create thumbnail maintaining aspect ratio
                    img.thumbnail(size, Image.Resampling.LANCZOS)
                    img.save(thumbnail_path, "JPEG", quality=85, optimize=True)
                    return str(thumbnail_path)
            
            elif extension == '.pdf':
                # PDF thumbnail
                doc = fitz.open(file_path)
                if len(doc) > 0:
                    page = doc[0]
                    mat = fitz.Matrix(2.0, 2.0)  # Higher resolution
                    pix = page.get_pixmap(matrix=mat)
                    img_data = pix.tobytes("ppm")
                    
                    with Image.open(io.BytesIO(img_data)) as img:
                        img.thumbnail(size, Image.Resampling.LANCZOS)
                        img.save(thumbnail_path, "JPEG", quality=85, optimize=True)
                    
                    doc.close()
                    return str(thumbnail_path)
            
            elif extension in ['.docx']:
                # Word document - extract first image or create text preview
                try:
                    doc = docx.Document(file_path)
                    
                    # Try to find first image
                    for rel in doc.part.rels.values():
                        if "image" in rel.target_ref:
                            image_data = rel.target_part.blob
                            with Image.open(io.BytesIO(image_data)) as img:
                                if img.mode in ('RGBA', 'LA', 'P'):
                                    img = img.convert('RGB')
                                img.thumbnail(size, Image.Resampling.LANCZOS)
                                img.save(thumbnail_path, "JPEG", quality=85)
                                return str(thumbnail_path)
                    
                    # If no image, create text preview
                    text = ""
                    for paragraph in doc.paragraphs[:5]:  # First 5 paragraphs
                        text += paragraph.text + "\n"
                    
                    if text.strip():
                        return await self.create_text_thumbnail(text[:200], thumbnail_path, size)
                        
                except Exception as e:
                    print(f"Error processing DOCX: {e}")
            
        except Exception as e:
            print(f"Error generating thumbnail for {file_path}: {e}")
        
        return None
    
    async def create_text_thumbnail(self, text: str, output_path: Path, size: tuple = (300, 300)) -> str:
        """Create a thumbnail image from text content"""
        try:
            # Create a simple text preview image
            img = Image.new('RGB', size, color='white')
            
            # This would require PIL with text rendering capabilities
            # For now, return a placeholder
            img.save(output_path, "JPEG", quality=85)
            return str(output_path)
            
        except Exception as e:
            print(f"Error creating text thumbnail: {e}")
            return None
    
    async def extract_text_content(self, file_path: Path) -> Optional[str]:
        """Extract text content for search indexing"""
        try:
            extension = file_path.suffix.lower()
            
            if extension == '.pdf':
                doc = fitz.open(file_path)
                text = ""
                for page in doc:
                    text += page.get_text() + "\n"
                doc.close()
                return text[:10000]  # Limit to 10KB of text
            
            elif extension in ['.txt', '.md', '.py', '.js', '.html', '.css']:
                async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = await f.read(10000)
                    return content
            
            elif extension == '.docx':
                doc = docx.Document(file_path)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                return text[:10000]
            
        except Exception as e:
            print(f"Error extracting text from {file_path}: {e}")
        
        return None
    
    async def create_preview(self, file_path: Path, document_id: str) -> Optional[str]:
        """Create preview for supported file types"""
        try:
            preview_path = self.previews_dir / f"{document_id}_preview.jpg"
            extension = file_path.suffix.lower()
            
            if extension == '.pdf':
                # PDF preview - first page at higher resolution
                doc = fitz.open(file_path)
                if len(doc) > 0:
                    page = doc[0]
                    mat = fitz.Matrix(3.0, 3.0)  # High resolution
                    pix = page.get_pixmap(matrix=mat)
                    img_data = pix.tobytes("ppm")
                    
                    with Image.open(io.BytesIO(img_data)) as img:
                        # Resize to max 800px width while maintaining aspect ratio
                        img.thumbnail((800, 1200), Image.Resampling.LANCZOS)
                        img.save(preview_path, "JPEG", quality=90, optimize=True)
                    
                    doc.close()
                    return str(preview_path)
            
            elif extension in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.webp']:
                # Image preview
                with Image.open(file_path) as img:
                    if img.mode in ('RGBA', 'LA', 'P'):
                        img = img.convert('RGB')
                    
                    # Create high-quality preview
                    img.thumbnail((800, 800), Image.Resampling.LANCZOS)
                    img.save(preview_path, "JPEG", quality=90, optimize=True)
                    return str(preview_path)
            
        except Exception as e:
            print(f"Error creating preview for {file_path}: {e}")
        
        return None
    
    def organize_file_by_date(self, file_path: Path, document_id: str, upload_date: datetime) -> Path:
        """Organize files by year/month structure"""
        year_month = upload_date.strftime("%Y/%m")
        organized_dir = self.documents_dir / year_month
        organized_dir.mkdir(parents=True, exist_ok=True)
        
        new_path = organized_dir / f"{document_id}{file_path.suffix}"
        
        # Move file to organized location
        shutil.move(str(file_path), str(new_path))
        return new_path
    
    def cleanup_temp_files(self, older_than_hours: int = 24):
        """Clean up temporary files older than specified hours"""
        try:
            cutoff_time = datetime.now().timestamp() - (older_than_hours * 3600)
            
            for file_path in self.temp_dir.iterdir():
                if file_path.is_file() and file_path.stat().st_mtime < cutoff_time:
                    file_path.unlink()
                    
        except Exception as e:
            print(f"Error cleaning up temp files: {e}")
    
    def get_storage_stats(self) -> Dict[str, Any]:
        """Get storage statistics"""
        try:
            total_size = 0
            file_count = 0
            
            for file_path in self.documents_dir.rglob("*"):
                if file_path.is_file():
                    total_size += file_path.stat().st_size
                    file_count += 1
            
            return {
                "total_files": file_count,
                "total_size_bytes": total_size,
                "total_size_mb": round(total_size / (1024 * 1024), 2),
                "total_size_gb": round(total_size / (1024 * 1024 * 1024), 2)
            }
        except Exception as e:
            print(f"Error getting storage stats: {e}")
            return {"error": str(e)}

# Global file manager instance
file_manager = FileManager()

# Import missing modules at the top
import io
